from models import (Campaign,
                    Event,
                    Result,
                    get_result_event_to_export)
import jinja2
from io import BytesIO
from xhtml2pdf import pisa
import os
from docx import Document
from stat_ import format_time
from docx.shared import Inches
# import pandas as pd
TEMPLATE_DIR = os.path.join(os.path.dirname(__file__), "export_templates")
environment = jinja2.Environment(loader=jinja2.FileSystemLoader(TEMPLATE_DIR))


input_dict = {'results': [{'email': 'a@b.c', 'open': True, 'click': True,
                                    'submit': False, 'report': False, 'os': 'windows', 'browser': 'chrome'},
                          {'email': 'a@b.c', 'open': True, 'click': True,
                           'submit': False, 'report': False, 'os': 'windows', 'browser': 'chrome'}
                          ],
              'details': [
    {'firstname': 'john', 'lastname': 'carter', 'ip': '0.0.0.0', 'email': 'john.c@c.co', 'location': 'test,test,test',
        'browser': 'firefox', 'os': 'linux', 'send': '01/01/23 00:00:00', 'open': '01/01/23 00:00:00', 'click': '01/01/23 00:00:00', 'submit': '01/01/23 00:00:00'},
],
    'browers': [{'name': 'chrome', 'count': 10}, {'name': 'firefox', 'count': 5}],
    'oses': [{'name': 'windows', 'count': 10}, {'name': 'ubuntu', 'count': 5}],
    'ips': [{'name': '0.0.0.0', 'count': 10}, {'name': '127.0.0.1', 'count': 5}],
    'campaign': {'name': 'test', 'status': 'complete', 'create_at': '01/01/23 00:00:00', 'launch_date': '01/01/23 00:00:00', 'complete_date': '03/03/23'},
    'envelop_sender': 'es', 'base_url': 'ur', 're_url': 're', 'capture_cred': False, 'capture_pass': False,
    'targets': 10, 'opened': 5, 'clicked': 3, 'submitted': 1}


def convert_html_to_pdf(data: dict):
    output = BytesIO()
    temp = environment.get_template('pdf.html')
    rendered = temp.render(data)
    # convert HTML to PDF
    pisa_status = pisa.CreatePDF(
        rendered,                # the HTML to convert
        dest=output)           # ByteIo handle to recieve result

    if pisa_status.err:
        raise Exception('cannot create pdf')
    # return bytes of pdf file
    return output.getvalue()


def export_pdf(campaign: Campaign):
    res = get_result_event_to_export(
        campaign_id=campaign.id, org_id=campaign.org_id)
    return convert_html_to_pdf(res)


# def convert_html_to_docx(data: dict):
#     document = Document()
#     document.add_heading('Executive Summary', 0)

#     # Add content to the document
#     document.add_paragraph(f"Campaign Results For: {data['campaign']['name']}")
#     document.add_paragraph(f"Status: {data['campaign']['status']}")
#     document.add_paragraph(f"Create: {data['campaign']['created_date']}")
#     document.add_paragraph(f"Started: {data['campaign']['launch_date']}")
#     document.add_paragraph(f"Completed: {data['campaign']['completed_date']}")

#     # Add more content...

#     # Save the document to a BytesIO object
#     output = BytesIO()
#     document.save(output)
#     output.seek(0)
#     return output
def convert_html_to_docx(data: dict):
    # results, timelines, statistics = static

    document = Document()
    document.add_heading('Executive Summary', 0)

    # Add content to the document
    document.add_heading(
        f"Campaign Results For: {data['campaign']['name']}", 1)
    document.add_paragraph(f"Status: {data['campaign']['status']}")
    document.add_paragraph(f"Create: {data['campaign']['created_date']}")
    document.add_paragraph(f"Started: {data['campaign']['launch_date']}")
    document.add_paragraph(f"Completed: {data['campaign']['completed_date']}")

    document.add_heading('Campaign Details', 1)
    document.add_paragraph(f"From: {data['envelope_sender']}")
    document.add_paragraph(f"Phish URL: {data['base_url']}")
    document.add_paragraph(f"Redirect URL: {data['envelope_sender']}")
    document.add_paragraph("Attachment: None Used")
    document.add_paragraph(f"Captured Credentials: {data['capture_cred']}")
    document.add_paragraph(f"Stored Passwords: {data['capture_pass']}")

    document.add_heading("High Level Results", 1)
    document.add_paragraph(f"Total Targets: {data['targets']}")
    document.add_paragraph(f"Total Email sent: {data['sent']}")

    document.add_heading(
        "The following totals indicate how many targets participated in each event type:", 1)
    document.add_paragraph(f"Individuals Who Opened Email: {data['opened']}")
    document.add_paragraph(f"Individuals Who Clicked: {data['clicked']}")
    document.add_paragraph(f"Individuals Who Submitted: {data['submitted']}")
    document.add_paragraph(f"Individuals Who Reported:: {data['reported']}")

    document.add_page_break()  # This will insert a page break
    document.add_heading('Summary of Events', 0)
    document.add_paragraph(
        "The following table summarizes who opened and clicked on emails sent in this campaign.")

    # Add campaign results table
    table = document.add_table(rows=1, cols=7)
    table.style = 'Table Grid'
    table.autofit = False

    # Set column widths
    col_widths = (Inches(2.2), Inches(0.1), Inches(0.1), Inches(
        0.1), Inches(0.1), Inches(1.3), Inches(1.3))
    for i, width in enumerate(col_widths):
        for row in table.rows:
            row.cells[i].width = width

    # Add header row
    header_cells = table.rows[0].cells
    header_cells[0].text = 'Email Address'
    header_cells[1].text = 'Open'
    header_cells[2].text = 'Click'
    header_cells[3].text = 'Submit'
    header_cells[4].text = 'Report'
    header_cells[5].text = 'OS'
    header_cells[6].text = 'Browser'

    # Add data rows
    for res in data['details']:
        row = table.add_row().cells
        row[0].text = res['email']
        row[1].text = '✔︎' if res['open'] else '✘'
        row[2].text = '✔︎' if res['click'] else '✘'
        row[3].text = '✔︎' if res['submit'] else '✘'
        row[4].text = '✔︎' if res['report'] else '✘'
        row[5].text = res['os'] if res['os'] else 'N/A'
        row[6].text = res['browser'] if res['browser'] else 'N/A'

    document.add_page_break()  # This will insert a page break

#### Detailed Findings ########

    for target in data['details']:
        document.add_heading('Detailed Findings', level=0)
        if target['open'] or target['click'] or target['submit']:
            document.add_heading(
                f"{target['firstname']} {target['lastname']}", level=2)
            document.add_paragraph(
                f"Department: {target['department'] if target.get('department') else '' }")
            document.add_paragraph(f"Email: {target['email']}")
            document.add_paragraph(f"Email sent: {target['send']}")

            if target['open']:
                document.add_paragraph("Email Previews")
                table = document.add_table(rows=1, cols=1)
                table.style = 'Table Grid'
                table.autofit = False
                # Set column widths
                col_widths = (Inches(1.8), Inches(1), Inches(1), Inches(
                    1.3), Inches(1.3))
                for i, width in enumerate(col_widths):
                    for row in table.rows:
                        row.cells[i].width = width
                row = table.rows[0].cells
                row[0].text = target['open']

            if target['click']:
                document.add_paragraph("Email Link Clicked")
                table = document.add_table(rows=1, cols=5)
                table.style = 'Table Grid'
                table.autofit = False
                # Set column widths
                col_widths = (Inches(1.8), Inches(1), Inches(1), Inches(
                    1.3), Inches(1.3))
                for i, width in enumerate(col_widths):
                    for row in table.rows:
                        row.cells[i].width = width

                header_cells = table.rows[0].cells
                header_cells[0].text = 'Time'
                header_cells[1].text = 'IP'
                header_cells[2].text = 'Location'
                header_cells[3].text = 'Browser'
                header_cells[4].text = 'Operating System'

                row = table.add_row().cells
                row[0].text = target['click'].strftime(
                    '%Y-%m-%d %H:%M:%S')  # Convert datetime to string
                row[1].text = target['ip']
                row[2].text = target['location']
                row[3].text = target['browser']
                row[4].text = target['os']

            if target['submit']:
                document.add_paragraph("Data Captured")
                table = document.add_table(rows=1, cols=6)
                table.style = 'Table Grid'
                table.autofit = False
                # Set column widths
                col_widths = (Inches(1.8), Inches(1), Inches(1), Inches(
                    1.3), Inches(1.3), Inches(1.3))
                for i, width in enumerate(col_widths):
                    for row in table.rows:
                        row.cells[i].width = width

                header_cells = table.rows[0].cells
                header_cells[0].text = 'Time'
                header_cells[1].text = 'IP'
                header_cells[2].text = 'Location'
                header_cells[3].text = 'Browser'
                header_cells[4].text = 'Operating System'
                header_cells[5].text = 'Data Captured'

                row = table.add_row().cells
                row[0].text = target['submit'].strftime(
                    '%Y-%m-%d %H:%M:%S')  # Convert datetime to string
                row[1].text = target['ip']
                row[2].text = target['location']
                row[3].text = target['browser']
                row[4].text = target['os']
                row[5].text = target['playload']
            document.add_page_break()

    # This will insert a page break
    # Add a section for statistics
################## Statistics ###################
    document.add_heading('Statistics', level=0)

    document.add_heading('Statistics Summary', level=2)

    # สร้างตาราง
    table = document.add_table(rows=1, cols=3)
    table.style = 'Table Grid'

    # กำหนดหัวข้อคอลัมน์
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Statistic'
    hdr_cells[1].text = 'Mean'
    hdr_cells[2].text = 'Standard Deviation'

    for statistics_a in data['statistics']:
        statistics_data = [
            ('Risk Percentage',
             statistics_a['mean_risk_percentage'], statistics_a['std_risk_percentage']),
            ('Time Sent To Submit',
             (statistics_a['mean_time_sent_to_submit']), (statistics_a['std_time_sent_to_submit'])),
            ('Time Sent To Open',
             (statistics_a['mean_time_sent_to_open']), (statistics_a['std_time_sent_to_open'])),
            ('Time Open To Click',
             (statistics_a['mean_time_open_to_click']), (statistics_a['std_time_open_to_click'])),
            ('Time Click To Submit',
             (statistics_a['mean_time_click_to_submit']), (statistics_a['std_time_click_to_submit'])),
            ('Time Sent To Report',
             (statistics_a['mean_time_sent_to_report']), (statistics_a['std_time_sent_to_report']))
        ]

    # เพิ่มข้อมูลลงในตาราง
    for statistic, mean, std_dev in statistics_data:
        row_cells = table.add_row().cells
        row_cells[0].text = statistic
        row_cells[1].text = str(mean)
        row_cells[2].text = str(std_dev)

    # Add statistics about browsers
    document.add_heading(
        'The following table shows the browsers seen:', level=2)
    browser_table = document.add_table(rows=1, cols=2)
    browser_table.style = 'Table Grid'
    browser_table.autofit = False
    browser_header_cells = browser_table.rows[0].cells
    browser_header_cells[0].text = 'Browser'
    browser_header_cells[1].text = 'Seen'

    for browser in data['browers']:
        row = browser_table.add_row().cells
        row[0].text = browser['name']
        row[1].text = str(browser['count'])

    # Add statistics about operating systems
    document.add_heading(
        'The following table shows the operating systems seen:', level=2)
    os_table = document.add_table(rows=1, cols=2)
    os_table.style = 'Table Grid'
    os_table.autofit = False
    os_header_cells = os_table.rows[0].cells
    os_header_cells[0].text = 'Operating System'
    os_header_cells[1].text = 'Seen'

    for os in data['oses']:
        row = os_table.add_row().cells
        row[0].text = os['name']
        row[1].text = str(os['count'])

    # Add statistics about IP addresses
    document.add_heading(
        'The following table shows the IP addresses captured:', level=2)
    ip_table = document.add_table(rows=1, cols=2)
    ip_table.style = 'Table Grid'
    ip_table.autofit = False
    ip_header_cells = ip_table.rows[0].cells
    ip_header_cells[0].text = 'IP Address'
    ip_header_cells[1].text = 'Seen'

    for ip in data['ips']:
        row = ip_table.add_row().cells
        row[0].text = ip['name']
        row[1].text = str(ip['count'])
    sections = document.sections
    for section in sections:
        section.page_width = Inches(10.5)  # ปรับตามความต้องการของคุณ

    output = BytesIO()
    document.save(output)

    return output.getvalue()


def export_docx(campaign: Campaign):
    res = get_result_event_to_export(
        campaign_id=campaign.id, org_id=campaign.org_id)
    return convert_html_to_docx(res)
